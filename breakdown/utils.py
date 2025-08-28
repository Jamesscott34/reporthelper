"""
Utility functions for the breakdown app.

This module contains helper functions for text extraction and file
processing, including pointer maps that link extracted text back to
original sources (pages, paragraphs, offsets).
"""

import os
import PyPDF2
from docx import Document as DocxDocument
import logging
import subprocess
import tempfile
import shutil
import zipfile
from typing import Tuple, Dict, Any, List

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """
    Extract text from a file based on its type.

    Args:
        file_path: Path to the file
        file_type: Type of the file (pdf, docx, doc, txt)

    Returns:
        Extracted text content
    """
    try:
        print(f"Extracting text from {file_path} (type: {file_type})")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        if file_type.lower() == 'pdf':
            return extract_text_from_pdf(file_path)
        if file_type.lower() == 'docx':
            return extract_text_from_docx(file_path)
        if file_type.lower() == 'doc':
            return extract_text_from_doc(file_path)
        if file_type.lower() == 'txt':
            return extract_text_from_txt(file_path)
        raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        print(f"Error extracting text from {file_path}: {e}")
        raise


def extract_text_with_pointers(
    file_path: str, file_type: str
) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text along with a structured pointer map to the original.

    Parameters
    ----------
    file_path : str
        Path to the file.
    file_type : str
        One of 'pdf', 'docx', 'doc', 'txt'.

    Returns
    -------
    (text, extraction_map) : Tuple[str, Dict[str, Any]]
        Extracted plain text and a JSON-serializable map describing
        where each piece of text came from (pages/paragraphs/offsets).
    """
    ft = file_type.lower()
    if ft == 'pdf':
        return _extract_pdf_with_pointers(file_path)
    if ft == 'docx':
        return _extract_docx_with_pointers(file_path)
    if ft == 'doc':
        # Convert to DOCX and reuse logic
        text = extract_text_from_doc(file_path)
        # Provide a coarse map based on line offsets
        lines = text.splitlines()
        pointer_lines: List[Dict[str, int]] = []
        cur = 0
        for i, ln in enumerate(lines):
            start = cur
            end = start + len(ln)
            pointer_lines.append({
                'line': i + 1,
                'char_start': start,
                'char_end': end,
            })
            cur = end + 1  # account for newline
        return text, {'type': 'doc', 'lines': pointer_lines}
    if ft == 'txt':
        return _extract_txt_with_pointers(file_path)
    raise ValueError(f"Unsupported file type for pointers: {file_type}")


def _extract_pdf_with_pointers(
    file_path: str,
) -> Tuple[str, Dict[str, Any]]:
    """Extract PDF text with page and line pointers.

    Returns
    -------
    Tuple[str, dict]
        Text and map: { 'type': 'pdf', 'pages': [
          { 'page': n,
            'lines': [ {'index': i, 'char_start': s, 'char_end': e} ] }
        ] }
    """
    text_parts: List[str] = []
    pages_map: List[Dict[str, Any]] = []
    cur = 0
    with open(file_path, 'rb') as fh:
        pdf = PyPDF2.PdfReader(fh)
        for pi, page in enumerate(pdf.pages):
            header = f"\n--- Page {pi + 1} ---\n"
            text_parts.append(header)
            cur += len(header)
            page_text = page.extract_text() or ''
            lines = page_text.splitlines()
            line_ptrs: List[Dict[str, int]] = []
            for li, ln in enumerate(lines):
                start = cur
                text_parts.append(ln + "\n")
                end = start + len(ln)
                cur = end + 1
                line_ptrs.append({
                    'index': li + 1,
                    'char_start': start,
                    'char_end': end,
                })
            pages_map.append({'page': pi + 1, 'lines': line_ptrs})
    full_text = ''.join(text_parts).strip()
    return full_text, {'type': 'pdf', 'pages': pages_map}


def _extract_docx_with_pointers(
    file_path: str,
) -> Tuple[str, Dict[str, Any]]:
    """Extract DOCX text with paragraph indices.

    Returns
    -------
    Tuple[str, dict]
        Text and map: { 'type': 'docx', 'paragraphs': [
          {'index': i, 'char_start': s, 'char_end': e}
        ]}
    """
    doc = DocxDocument(file_path)
    text_parts: List[str] = []
    ptrs: List[Dict[str, int]] = []
    cur = 0
    para_index = 0
    # Paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            para_index += 1
            start = cur
            text_parts.append(paragraph.text + "\n")
            end = start + len(paragraph.text)
            cur = end + 1
            ptrs.append({
                'index': para_index,
                'char_start': start,
                'char_end': end,
            })
    # Tables (optional): append as additional paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    para_index += 1
                    start = cur
                    text_parts.append(cell_text + "\n")
                    end = start + len(cell_text)
                    cur = end + 1
                    ptrs.append({
                        'index': para_index,
                        'char_start': start,
                        'char_end': end,
                    })
    full_text = ''.join(text_parts).strip()
    return full_text, {'type': 'docx', 'paragraphs': ptrs}


def _extract_txt_with_pointers(
    file_path: str,
) -> Tuple[str, Dict[str, Any]]:
    """Extract TXT with character offsets and line indices."""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    text = ''
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as fh:
                text = fh.read()
                break
        except UnicodeDecodeError:
            continue
    if not text.strip():
        raise Exception('No text could be extracted from TXT')
    lines = text.splitlines()
    ptrs: List[Dict[str, int]] = []
    cur = 0
    out_parts: List[str] = []
    for i, ln in enumerate(lines):
        start = cur
        out_parts.append(ln + "\n")
        end = start + len(ln)
        cur = end + 1
        ptrs.append({
            'line': i + 1,
            'char_start': start,
            'char_end': end,
        })
    return ''.join(out_parts).strip(), {'type': 'txt', 'lines': ptrs}


def unpack_zip_to_temp(zip_file_path: str) -> str:
    """Unpack a ZIP to a temporary directory and return its path."""
    tmp_dir = tempfile.mkdtemp()
    with zipfile.ZipFile(zip_file_path, 'r') as zf:
        zf.extractall(tmp_dir)
    return tmp_dir


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content
    """
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            print(f"PDF has {len(pdf_reader.pages)} pages")
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {i+1} ---\n{page_text}\n"
                else:
                    print(f"Warning: No text extracted from page {i+1}")
        text = text.strip()
        if not text:
            raise Exception("No text could be extracted from PDF")
        print(f"Successfully extracted {len(text)} characters from PDF")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        raise


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content
    """
    try:
        doc = DocxDocument(file_path)
        text = ""
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        text = text.strip()
        if not text:
            raise Exception("No text could be extracted from DOCX")
        print(
            "Successfully extracted "
            f"{len(text)} characters from DOCX"
        )
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        raise


def extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from a DOC file using LibreOffice conversion.

    Args:
        file_path: Path to the DOC file

    Returns:
        Extracted text content
    """
    try:
        # Create a temporary directory for conversion
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, "converted.docx")
        # Convert DOC to DOCX using LibreOffice
        cmd = [
            'libreoffice', '--headless', '--convert-to', 'docx',
            '--outdir', temp_dir, file_path,
        ]
        print(
            "Converting DOC to DOCX using command: " + ' '.join(cmd)
        )
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            raise Exception(
                f"LibreOffice conversion failed: {result.stderr}"
            )
        # Check if conversion was successful
        if not os.path.exists(temp_docx):
            converted = [
                f for f in os.listdir(temp_dir) if f.endswith('.docx')
            ]
            if converted:
                temp_docx = os.path.join(temp_dir, converted[0])
            else:
                raise Exception(
                    "DOC to DOCX conversion failed - no output file found"
                )
        # Extract text from the converted DOCX
        text = extract_text_from_docx(temp_docx)
        # Clean up temporary files
        shutil.rmtree(temp_dir, ignore_errors=True)
        return text
    except subprocess.TimeoutExpired:
        raise Exception("DOC conversion timed out")
    except Exception as e:
        logger.error(f"Error extracting text from DOC {file_path}: {e}")
        if 'temp_dir' in locals():
            shutil.rmtree(temp_dir, ignore_errors=True)
        raise


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from a TXT file.

    Args:
        file_path: Path to the TXT file

    Returns:
        Extracted text content
    """
    try:
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    if text.strip():
                        print(
                            "Successfully extracted "
                            f"{len(text)} characters from TXT using "
                            f"{encoding} encoding"
                        )
                        return text.strip()
            except UnicodeDecodeError:
                continue
        raise Exception(
            "Could not decode TXT file with any supported encoding"
        )
    except Exception as e:
        logger.error(f"Error extracting text from TXT {file_path}: {e}")
        raise


def validate_file_type(file_name: str) -> bool:
    """
    Validate if the file type is supported.

    Args:
        file_name: Name of the file

    Returns:
        True if file type is supported, False otherwise
    """
    allowed_extensions = ['.pdf', '.docx', '.doc', '.txt']
    file_extension = os.path.splitext(file_name)[1].lower()
    return file_extension in allowed_extensions


def get_file_size_mb(file_path: str) -> float:
    """
    Get file size in megabytes.

    Args:
        file_path: Path to the file

    Returns:
        File size in MB
    """
    return os.path.getsize(file_path) / (1024 * 1024)
